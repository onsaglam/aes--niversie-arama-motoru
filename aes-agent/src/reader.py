"""
reader.py — Öğrenci profil Word dosyası, transkript PDF ve dil belgesi PDF okuma.
"""
import os
import re
import base64
import json
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import anthropic
    CLAUDE_OK = True
except ImportError:
    CLAUDE_OK = False


# ─── GPA Dönüşüm ─────────────────────────────────────────────────────────────

def _tr_100_to_40(grade_100: float) -> float:
    """Türkiye 100 puan skalasını 4.0 skalasına çevir."""
    if grade_100 >= 90: return 4.0
    if grade_100 >= 85: return 3.7
    if grade_100 >= 80: return 3.3
    if grade_100 >= 75: return 3.0
    if grade_100 >= 70: return 2.7
    if grade_100 >= 65: return 2.3
    if grade_100 >= 60: return 2.0
    if grade_100 >= 55: return 1.7
    if grade_100 >= 50: return 1.3
    return 1.0


def convert_to_german_scale(raw: str) -> Optional[float]:
    """
    Her türlü GPA string'ini Almanya 1.0-5.0 skalasına çevir.
    Almanya: 1.0=Sehr gut, 2.0=Gut, 3.0=Befriedigend, 4.0=Ausreichend
    """
    if not raw:
        return None
    raw = raw.strip().replace(",", ".")

    # "3.2/4.0" veya "3.2 / 4.0" formatı
    m = re.match(r"([\d.]+)\s*/\s*([\d.]+)", raw)
    if m:
        try:
            value, scale = float(m.group(1)), float(m.group(2))
            if scale == 4.0:
                # Bavyera formülü
                return round(1.0 + 3.0 * (4.0 - value) / (4.0 - 1.0), 1)
            if scale == 100 or scale == 100.0:
                gpa_40 = _tr_100_to_40(value)
                return round(1.0 + 3.0 * (4.0 - gpa_40) / (4.0 - 1.0), 1)
        except (ValueError, ZeroDivisionError):
            pass

    # Zaten Almanya skalasında: "2.5" veya "2.5 (DE)"
    m2 = re.match(r"([\d.]+)\s*(?:\(DE\)|\(de\))?$", raw)
    if m2:
        try:
            val = float(m2.group(1))
            if 1.0 <= val <= 5.0:
                return val
        except ValueError:
            pass

    return None


# ─── StudentProfile ───────────────────────────────────────────────────────────

@dataclass
class StudentProfile:
    # Kişisel
    name: str = ""
    email: str = ""
    phone: str = ""
    nationality: str = "Türk"

    # Eğitim
    current_university: str = ""
    department: str = ""
    gpa_turkish: str = ""
    gpa_german: str = ""         # Almanya skalasına dönüştürülmüş (string)
    gpa_german_float: Optional[float] = None  # Karşılaştırma için float
    graduation_date: str = ""
    diploma_status: str = ""

    # Dil
    german_level: str = ""
    english_level: str = ""

    # Hedef
    desired_field: str = ""
    degree_type: str = "Master"
    program_language: str = ""
    preferred_cities: list = field(default_factory=list)
    start_semester: str = ""

    # Tercihler
    free_tuition_important: bool = True
    university_type: str = ""
    accept_nc: bool = True
    conditional_admission: bool = True
    part_time_important: bool = False

    # Belgeler
    has_transcript: bool = False
    has_language_cert: bool = False
    has_motivation_letter: bool = False
    has_cv: bool = False

    # PDF'ten zenginleştirilmiş alanlar
    transcript_data: dict = field(default_factory=dict)
    language_cert_data: dict = field(default_factory=dict)

    # Danışman notları
    advisor_notes: str = ""

    def to_summary(self) -> str:
        gpa_display = self.gpa_german or self.gpa_turkish or "belirtilmedi"
        lines = [
            f"Ad: {self.name}",
            f"Alan: {self.desired_field} ({self.degree_type})",
            f"GPA: {gpa_display}",
            f"Almanca: {self.german_level or 'Yok'}",
            f"İngilizce: {self.english_level or 'Yok'}",
            f"Program dili: {self.program_language or 'Fark etmez'}",
            f"Şehir tercihi: {', '.join(self.preferred_cities) if self.preferred_cities else 'Fark etmez'}",
            f"Dönem: {self.start_semester or 'belirtilmedi'}",
            f"Şartlı kabul: {'Evet' if self.conditional_admission else 'Hayır'}",
        ]
        if self.advisor_notes:
            lines.append(f"Notlar: {self.advisor_notes}")
        return "\n".join(lines)


# ─── Ana okuma fonksiyonu ─────────────────────────────────────────────────────

def read_profile(student_folder: str | Path) -> StudentProfile:
    """Öğrenci klasöründeki tüm dosyaları oku ve StudentProfile döndür."""
    folder = Path(student_folder)
    profile = StudentProfile()

    # 1. Dashboard'dan kaydedilen JSON profil (öncelikli — en güncel veri)
    json_path = folder / "profil.json"
    if json_path.exists():
        try:
            profile = _read_json(json_path)
        except Exception:
            pass

    # 2. Word profil dosyası (JSON yoksa veya JSON'da eksik alanlar için)
    docx_path = folder / "profil.docx"
    if docx_path.exists() and Document and not json_path.exists():
        profile = _read_docx(docx_path)
    elif not docx_path.exists() and not json_path.exists():
        pass  # Profil dosyası yoksa klasör adını isim olarak kullan

    # 2. Belge varlık kontrolü
    profile.has_transcript       = (folder / "transkript.pdf").exists()
    profile.has_language_cert    = (folder / "dil_belgesi.pdf").exists()
    profile.has_motivation_letter= (folder / "motivasyon.docx").exists()
    profile.has_cv               = (folder / "cv.pdf").exists()

    # 3. Transkript PDF — Claude ile
    if profile.has_transcript and CLAUDE_OK:
        try:
            t_data = _extract_transcript(folder / "transkript.pdf")
            profile.transcript_data = t_data
            # GPA'yı PDF'ten zenginleştir (profil.docx boşsa)
            if not profile.gpa_turkish and t_data.get("gpa_original"):
                profile.gpa_turkish = t_data["gpa_original"]
            if not profile.current_university and t_data.get("university"):
                profile.current_university = t_data["university"]
            if not profile.department and t_data.get("department"):
                profile.department = t_data["department"]
        except Exception:
            pass  # PDF okunamazsa devam et

    # 4. Dil belgesi PDF — Claude ile
    if profile.has_language_cert and CLAUDE_OK:
        try:
            l_data = _extract_language_cert(folder / "dil_belgesi.pdf")
            profile.language_cert_data = l_data
            normalized = l_data.get("normalized", "")
            lang = l_data.get("language", "")
            if normalized:
                if "Almanca" in lang or "German" in lang or "Deutsch" in lang:
                    if not profile.german_level:
                        profile.german_level = normalized
                elif "İngilizce" in lang or "English" in lang:
                    if not profile.english_level:
                        profile.english_level = normalized
        except Exception:
            pass

    # 5. Program dili otomatik tespiti (danışman boş bırakmışsa sertifikadan türet)
    #    Kural: Sadece Almanca sertifika → Almanca program
    #           Sadece İngilizce sertifika → İngilizce program
    #           Her ikisi → profil.docx'teki tercih korunur
    if not profile.program_language or profile.program_language.lower() in ("fark etmez", "any", ""):
        has_german  = bool(profile.german_level  and profile.german_level.lower()  not in ("yok", "none", ""))
        has_english = bool(profile.english_level and profile.english_level.lower() not in ("yok", "none", ""))
        if has_german and not has_english:
            profile.program_language = "Almanca"
        elif has_english and not has_german:
            profile.program_language = "İngilizce"
        # İkisi de varsa veya ikisi de yoksa → boş bırak (her ikisini de ara)

    # 6. GPA Almanya skalasına dönüşüm
    gpa_raw = profile.gpa_german or profile.gpa_turkish or ""
    gpa_float = convert_to_german_scale(gpa_raw)
    if gpa_float:
        profile.gpa_german_float = gpa_float
        if not profile.gpa_german:
            profile.gpa_german = str(gpa_float)

    # 6. İsim fallback
    if not profile.name:
        profile.name = folder.name.replace("_", " ")

    return profile


# ─── JSON okuma (dashboard'dan kaydedilen profil) ────────────────────────────

def _read_json(path: Path) -> StudentProfile:
    """Dashboard'dan kaydedilen profil.json'u oku."""
    data = json.loads(path.read_text(encoding="utf-8"))
    profile = StudentProfile()
    for key, value in data.items():
        if hasattr(profile, key) and value is not None:
            if key == "preferred_cities":
                # String veya list kabul et
                if isinstance(value, str):
                    cities = [c.strip() for c in value.replace("،", ",").split(",") if c.strip()]
                    profile.preferred_cities = cities
                elif isinstance(value, list):
                    profile.preferred_cities = value
            elif key in ("free_tuition_important", "accept_nc", "conditional_admission"):
                profile.__dict__[key] = bool(value)
            else:
                setattr(profile, key, value)
    return profile


# ─── Word okuma ───────────────────────────────────────────────────────────────

def _read_docx(path: Path) -> StudentProfile:
    doc = Document(str(path))
    profile = StudentProfile()
    text = "\n".join(p.text for p in doc.paragraphs)

    def extract(pattern: str, default: str = "") -> str:
        m = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        return m.group(1).strip() if m else default

    profile.name               = extract(r"Ad Soyad[:\s]+\[?([^\]\n\[]+)")
    profile.email              = extract(r"E-posta[:\s]+\[?([^\]\n\[]+)")
    profile.current_university = extract(r"Mevcut.+Okul[:\s]+\[?([^\]\n\[]+)")
    profile.department         = extract(r"Bölüm[:\s]+\[?([^\]\n\[]+)")
    profile.gpa_turkish        = extract(r"Not Ortalamas[ıi][:\s]+\[?([^\]\n\[]+)")
    profile.graduation_date    = extract(r"Mezuniyet Tarihi[:\s]+\[?([^\]\n\[]+)")
    profile.german_level       = extract(r"Almanca[:\s]+\[?([^\]\n\[]+)")
    profile.english_level      = extract(r"İngilizce[:\s]+\[?([^\]\n\[]+)")
    profile.desired_field      = extract(r"İstenen Alan[:\s]+\[?([^\]\n\[]+)")
    profile.degree_type        = extract(r"Derece Türü[:\s]+\[?([^\]\n\[]+)", "Master")
    profile.program_language   = extract(r"Program Dili[:\s]+\[?([^\]\n\[]+)")
    profile.start_semester     = extract(r"Başlangıç Dönemi[:\s]+\[?([^\]\n\[]+)")
    profile.advisor_notes      = extract(r"EK BİLGİLER.*?\n(.+?)(?=\n■|\Z)", "")

    # Boş placeholder değerleri temizle
    placeholder_pattern = r"^\[.*\]$|^örn:.*$"
    for attr in ["name", "email", "current_university", "department",
                 "gpa_turkish", "german_level", "english_level",
                 "desired_field", "program_language", "start_semester"]:
        val = getattr(profile, attr)
        if val and re.match(placeholder_pattern, val.strip(), re.IGNORECASE):
            setattr(profile, attr, "")

    # Tercih şehirler
    cities_raw = extract(r"Tercih Şehirler[:\s]+\[?([^\]\n\[]+)")
    if cities_raw and "fark etmez" not in cities_raw.lower() and not re.match(placeholder_pattern, cities_raw):
        profile.preferred_cities = [c.strip() for c in re.split(r"[,،،]", cities_raw) if c.strip()]

    # Şartlı kabul
    cond_raw = extract(r"Şartlı Kabul[:\s]+\[?([^\]\n\[]+)", "Evet")
    profile.conditional_admission = "evet" in cond_raw.lower() or "yes" in cond_raw.lower()

    # GPA — Almanca skalasında belirtilmişse al
    gpa_raw = profile.gpa_turkish
    gpa_float = convert_to_german_scale(gpa_raw)
    if gpa_float:
        profile.gpa_german = str(gpa_float)
        profile.gpa_german_float = gpa_float

    return profile


# ─── Transkript PDF — Claude Vision ──────────────────────────────────────────

def _extract_transcript(pdf_path: Path) -> dict:
    """Transkript PDF'inden Claude ile veri çıkar."""
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY", ""))
    b64 = base64.standard_b64encode(pdf_path.read_bytes()).decode("utf-8")

    prompt = """Bu transkript belgesinden şunları çıkar.
SADECE JSON döndür, başka metin ekleme:
{
  "university": "üniversite adı",
  "department": "bölüm adı",
  "degree_level": "Bachelor/Master/Lisans/Yüksek Lisans",
  "gpa_original": "orijinal not (örn: 3.25/4.00 veya 2.85/4.0 veya 75.5/100)",
  "gpa_scale": "4.0 veya 100",
  "graduation_year": "yıl veya null",
  "status": "mezun/devam ediyor",
  "language": "eğitim dili (Türkçe/İngilizce vb.)"
}"""

    resp = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=500,
        messages=[{"role": "user", "content": [
            {"type": "document", "source": {
                "type": "base64",
                "media_type": "application/pdf",
                "data": b64
            }},
            {"type": "text", "text": prompt}
        ]}]
    )
    raw = resp.content[0].text.strip()
    m = re.search(r"\{[\s\S]*\}", raw)
    return json.loads(m.group()) if m else {}


# ─── Dil Belgesi PDF — Claude Vision ─────────────────────────────────────────

def _extract_language_cert(pdf_path: Path) -> dict:
    """Dil belgesi PDF'inden Claude ile sertifika türü ve seviyeyi çıkar."""
    client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY", ""))
    b64 = base64.standard_b64encode(pdf_path.read_bytes()).decode("utf-8")

    prompt = """Bu dil sertifikası belgesinden şunları çıkar.
SADECE JSON döndür:
{
  "cert_type": "TestDaF/DSH/Goethe/IELTS/TOEFL/ÖSD/telc/Cambridge veya diğer",
  "language": "Almanca veya İngilizce veya diğer",
  "level_or_score": "ham skor veya seviye (örn: 16 veya B2 veya 6.5)",
  "cefr_level": "A1/A2/B1/B2/C1/C2 veya null",
  "exam_date": "YYYY-MM veya null",
  "validity_expiry": "son geçerlilik tarihi veya 'ömür boyu' veya null"
}"""

    resp = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=400,
        messages=[{"role": "user", "content": [
            {"type": "document", "source": {
                "type": "base64",
                "media_type": "application/pdf",
                "data": b64
            }},
            {"type": "text", "text": prompt}
        ]}]
    )
    raw = resp.content[0].text.strip()
    m = re.search(r"\{[\s\S]*\}", raw)
    data = json.loads(m.group()) if m else {}

    # Normalize: ham skor → standart format string
    cert  = data.get("cert_type", "")
    score = str(data.get("level_or_score", ""))
    cefr  = data.get("cefr_level", "")

    if "TestDaF" in cert:
        data["normalized"] = f"TestDaF {score}"
    elif "DSH" in cert:
        data["normalized"] = f"DSH-{score.replace('DSH-','').replace('DSH','').strip()}"
    elif "IELTS" in cert:
        data["normalized"] = f"IELTS {score}"
    elif "TOEFL" in cert:
        data["normalized"] = f"TOEFL {score}"
    elif "Goethe" in cert:
        data["normalized"] = f"Goethe {cefr or score}"
    elif "telc" in cert:
        data["normalized"] = f"telc {cefr or score}"
    elif "ÖSD" in cert:
        data["normalized"] = f"ÖSD {cefr or score}"
    else:
        data["normalized"] = f"{cert} {score}".strip()

    return data


# ─── Şablon oluştur ───────────────────────────────────────────────────────────

def create_template(output_path: str | Path):
    """Boş profil şablonu oluştur."""
    if not Document:
        print("python-docx kurulu değil. pip install python-docx")
        return

    doc = Document()
    doc.add_heading("AES — Öğrenci Danışmanlık Profili", 0)

    sections = [
        ("KİŞİSEL BİLGİLER", [
            ("Ad Soyad", ""),
            ("Doğum Tarihi", ""),
            ("Milliyet", "Türk"),
            ("E-posta", ""),
            ("Telefon", ""),
        ]),
        ("EĞİTİM DURUMU", [
            ("Mevcut/Son Okul", ""),
            ("Bölüm", ""),
            ("Not Ortalaması", "örn: 2.8 (DE) veya 3.2/4.0 (TR)"),
            ("Mezuniyet Tarihi", "örn: Haziran 2024"),
            ("Diploma Durumu", "Alındı / Haziran 2025'te alınacak"),
        ]),
        ("DİL BELGELERİ", [
            ("Almanca", "TestDaF 16 / DSH-2 / Goethe C1 / Yok"),
            ("İngilizce", "IELTS 6.5 / TOEFL 88 / Yok"),
        ]),
        ("ALMANYA'DA HEDEF", [
            ("İstenen Alan", "örn: Makine Mühendisliği / Yapay Zeka"),
            ("Derece Türü", "Master / Bachelor / Ausbildung"),
            ("Program Dili", "Almanca / İngilizce / Fark etmez"),
            ("Tercih Şehirler", "örn: München, Berlin veya Fark etmez"),
            ("Başlangıç Dönemi", "WiSe 2025/26 / SoSe 2026"),
        ]),
        ("TERCİHLER", [
            ("Ücretsiz/Düşük Ücret", "Önemli / Fark etmez"),
            ("Üniversite Türü", "TU / FH / Fark etmez"),
            ("NC'li Bölüm", "Evet / Hayır"),
            ("Şartlı Kabul", "Evet kabul ediyorum / Hayır"),
            ("Part-time İmkanı", "Önemli / Fark etmez"),
        ]),
        ("EK BİLGİLER / DANIŞMAN NOTLARI", [
            ("Notlar", ""),
        ]),
    ]

    for section_title, fields in sections:
        doc.add_heading(f"■ {section_title}", level=1)
        for label, placeholder in fields:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}:  ")
            run.bold = True
            p.add_run(f"[{placeholder}]" if placeholder else "[ ]")

    doc.save(str(output_path))
    print(f"✅ Şablon oluşturuldu: {output_path}")
