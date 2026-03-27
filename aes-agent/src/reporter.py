"""
reporter.py — AES marka renkleriyle Word ve Excel raporu üret.
"""
from pathlib import Path
from datetime import datetime
from typing import List

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    XLSX_OK = True
except ImportError:
    XLSX_OK = False


# ─── AES Marka Renkleri ───────────────────────────────────────────────────────

AES_NAVY  = RGBColor(0x0A, 0x1F, 0x44) if DOCX_OK else None   # Koyu lacivert
AES_GOLD  = RGBColor(0xD4, 0xAF, 0x37) if DOCX_OK else None   # Altın
AES_GRAY  = RGBColor(0x80, 0x80, 0x80) if DOCX_OK else None   # Gri (footer)

COLOR_UYGUN     = RGBColor(0x16, 0xa3, 0x4a) if DOCX_OK else None  # Yeşil
COLOR_SARTLI    = RGBColor(0xca, 0x8a, 0x04) if DOCX_OK else None  # Sarı
COLOR_UYGUN_DEG = RGBColor(0xdc, 0x26, 0x26) if DOCX_OK else None  # Kırmızı

STATUS_INFO = {
    "uygun":       {"label": "✅ UYGUN",        "color": COLOR_UYGUN,     "fill_hex": "dcfce7"},
    "sartli":      {"label": "⚠️ ŞARTLI",       "color": COLOR_SARTLI,    "fill_hex": "fef9c3"},
    "uygun_degil": {"label": "❌ UYGUN DEĞİL",  "color": COLOR_UYGUN_DEG, "fill_hex": "fee2e2"},
}


# ─── Veri Kalite Kontrolü ─────────────────────────────────────────────────────

def validate_report_data(programs: list) -> list[str]:
    """Rapora girmeden önce veri kalitesini kontrol et."""
    warnings = []
    for p in programs:
        if not p.university:
            warnings.append(f"Boş üniversite adı: {p.url or 'URL yok'}")
        if not p.eligibility:
            warnings.append(f"Değerlendirme eksik: {p.university}")
        if p.confidence < 0.4:
            warnings.append(f"Düşük güvenilirlik ({p.confidence:.1f}): {p.university} — {p.program[:30]}")
    return warnings


# ─── Word Raporu ──────────────────────────────────────────────────────────────

def generate_word_report(programs: list, profile, output_dir: Path) -> Path | None:
    """AES marka renkleriyle Word raporu oluştur."""
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    if not DOCX_OK:
        print("⚠️  python-docx kurulu değil, Word raporu atlandı.")
        return None

    # Veri kalite kontrolü
    warnings = validate_report_data(programs)
    if warnings:
        print(f"⚠️  Rapor uyarıları ({len(warnings)} adet):")
        for w in warnings[:5]:
            print(f"   - {w}")

    doc = Document()

    # Sayfa kenar boşlukları
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Başlık ───────────────────────────────────────────────────────
    title = doc.add_heading("", 0)
    run   = title.add_run("Almanya Üniversite Araştırma Raporu")
    run.font.color.rgb = AES_NAVY
    run.font.size      = Pt(20)
    title.alignment    = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"Hazırlayan: AES — Almanya Eğitim Serüveni\n{datetime.now().strftime('%d.%m.%Y')}")
    sr.font.color.rgb = AES_GRAY
    sr.font.size      = Pt(11)
    doc.add_paragraph()

    # ── Öğrenci Profili ───────────────────────────────────────────────
    _add_heading(doc, "1. Öğrenci Profili")

    profile_rows = [
        ("Ad Soyad",         profile.name),
        ("Hedef Alan",        profile.desired_field),
        ("Derece Türü",       profile.degree_type),
        ("GPA (Orijinal)",    profile.gpa_turkish or "—"),
        ("GPA (Almanya)",     profile.gpa_german or "Hesaplanmadı"),
        ("Almanca",           profile.german_level or "Belirtilmedi"),
        ("İngilizce",         profile.english_level or "Belirtilmedi"),
        ("Program Dili",      profile.program_language or "Fark etmez"),
        ("Tercih Şehir",      ", ".join(profile.preferred_cities) if profile.preferred_cities else "Fark etmez"),
        ("Başlangıç Dönemi",  profile.start_semester or "—"),
        ("Şartlı Kabul",      "Kabul ediyor" if profile.conditional_admission else "Kabul etmiyor"),
    ]
    tbl = doc.add_table(rows=len(profile_rows), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(profile_rows):
        tbl.cell(i, 0).text = label
        tbl.cell(i, 1).text = str(value)
        run_ = tbl.cell(i, 0).paragraphs[0].runs
        if run_:
            run_[0].bold = True
        tbl.cell(i, 0).width = Cm(5)
        tbl.cell(i, 1).width = Cm(10)
    doc.add_paragraph()

    # ── Araştırma Özeti ───────────────────────────────────────────────
    _add_heading(doc, "2. Araştırma Özeti")

    uygun     = [p for p in programs if p.eligibility == "uygun"]
    sartli    = [p for p in programs if p.eligibility == "sartli"]
    uygun_deg = [p for p in programs if p.eligibility == "uygun_degil"]

    ozet = doc.add_paragraph()
    ozet.add_run(f"Toplam {len(programs)} program araştırıldı.  ").bold = True
    _colored_run(ozet, f"✅ {len(uygun)} uygun", COLOR_UYGUN)
    ozet.add_run("   ")
    _colored_run(ozet, f"⚠️ {len(sartli)} şartlı", COLOR_SARTLI)
    ozet.add_run("   ")
    _colored_run(ozet, f"❌ {len(uygun_deg)} uygun değil", COLOR_UYGUN_DEG)
    doc.add_paragraph()

    # ── Program Listeleri ─────────────────────────────────────────────
    groups = [
        (uygun,     "uygun",       3),
        (sartli,    "sartli",      4),
        (uygun_deg, "uygun_degil", 5),
    ]

    for group, status_key, section_num in groups:
        if not group:
            continue
        info = STATUS_INFO[status_key]
        h = doc.add_heading(f"{section_num}. {info['label']} Programlar ({len(group)})", level=1)
        if info["color"]:
            h.runs[0].font.color.rgb = info["color"]

        for prog in group:
            ph  = doc.add_paragraph()
            run = ph.add_run(f"{prog.university} — {prog.program}")
            run.bold           = True
            run.font.size      = Pt(12)
            run.font.color.rgb = AES_NAVY

            rows_data = [
                ("Şehir",           prog.city or "—"),
                ("Program Dili",    prog.language or "—"),
                ("WiSe Deadline",   prog.deadline_wise or "—"),
                ("SoSe Deadline",   prog.deadline_sose or "—"),
                ("Almanca Şartı",   prog.german_requirement or "—"),
                ("İngilizce Şartı", prog.english_requirement or "—"),
                ("NC Değeri",       prog.nc_value or "Zulassungsfrei"),
                ("uni-assist",      "Gerekli ✓" if prog.uni_assist_required else "Direkt başvuru"),
                ("Şartlı Kabul",    "Mevcut" if prog.conditional_admission else "Yok"),
                ("Değerlendirme",   prog.eligibility_reason or "—"),
            ]
            t = doc.add_table(rows=len(rows_data), cols=2)
            t.style = "Table Grid"
            for i, (k, v) in enumerate(rows_data):
                t.cell(i, 0).text = k
                t.cell(i, 1).text = str(v)
                runs_ = t.cell(i, 0).paragraphs[0].runs
                if runs_:
                    runs_[0].bold = True

            if prog.issues:
                ip = doc.add_paragraph()
                ir = ip.add_run("⚠️ Eksiklikler: " + " | ".join(prog.issues))
                ir.font.color.rgb = COLOR_SARTLI
                ir.font.size      = Pt(10)

            if prog.url:
                up = doc.add_paragraph()
                up.add_run("🔗 Başvuru: ").bold = True
                up.add_run(prog.url).font.color.rgb = RGBColor(0x00, 0x56, 0xCC)

            doc.add_paragraph()

    # ── Footer ────────────────────────────────────────────────────────
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        f"Bu rapor AES — Almanya Eğitim Serüveni tarafından oluşturulmuştur.\n"
        f"aes-kompass.com | Bremen, Germany | {datetime.now().strftime('%d.%m.%Y %H:%M')}\n"
        f"⚠️  Bilgiler araştırma tarihindeki durumu yansıtır. Başvuru öncesi üniversite sitelerini doğrulayınız."
    )
    fr.font.color.rgb = AES_GRAY
    fr.font.size      = Pt(9)

    out_path = output_dir / f"sonuc_raporu_{datetime.now().strftime('%Y%m%d')}.docx"
    doc.save(str(out_path))
    print(f"📄 Word raporu: {out_path}")
    return out_path


def _add_heading(doc, text: str):
    """AES lacivert rengiyle bölüm başlığı ekle."""
    h    = doc.add_heading(text, level=1)
    run_ = h.runs
    if run_:
        run_[0].font.color.rgb = AES_NAVY


def _colored_run(paragraph, text: str, color):
    """Renkli metin run'ı ekle."""
    r = paragraph.add_run(text)
    r.bold = True
    if color:
        r.font.color.rgb = color
    return r


# ─── Excel Raporu ─────────────────────────────────────────────────────────────

COLUMNS = [
    # (başlık, genişlik)
    ("Üniversite",      22),
    ("Şehir",           12),
    ("Program",         28),
    ("Dil",             12),
    ("WiSe Deadline",   14),
    ("SoSe Deadline",   14),
    ("Almanca Şartı",   14),
    ("İngilizce Şartı", 14),
    ("NC",              10),
    ("uni-assist",      10),
    ("Şartlı Kabul",    12),
    ("Uygunluk",        14),
    ("Değerlendirme",   42),
    ("Başvuru URL",     38),
]


def generate_excel_report(programs: list, profile, output_dir: Path) -> Path | None:
    """AES marka renkleriyle Excel raporu oluştur."""
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    if not XLSX_OK:
        print("⚠️  openpyxl kurulu değil, Excel raporu atlandı.")
        return None

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Üniversite Listesi"

    # Header stili — AES lacivert + altın yazı
    header_fill  = PatternFill("solid", fgColor="0A1F44")
    header_font  = Font(color="D4AF37", bold=True, size=10)
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    header_border = Border(bottom=Side(style="medium", color="D4AF37"))

    ws.row_dimensions[1].height = 30
    for col_idx, (header, width) in enumerate(COLUMNS, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill      = header_fill
        cell.font      = header_font
        cell.alignment = header_align
        cell.border    = header_border
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    # Renk dolgular
    status_fills = {
        "uygun":       PatternFill("solid", fgColor="dcfce7"),
        "sartli":      PatternFill("solid", fgColor="fef9c3"),
        "uygun_degil": PatternFill("solid", fgColor="fee2e2"),
    }
    status_labels = {
        "uygun":       "✅ Uygun",
        "sartli":      "⚠️ Şartlı",
        "uygun_degil": "❌ Uygun Değil",
    }

    # Önce uygun, sonra şartlı, sonra uygun değil
    order = {"uygun": 0, "sartli": 1, "uygun_degil": 2}
    sorted_programs = sorted(programs, key=lambda p: order.get(p.eligibility, 3))

    for row_idx, prog in enumerate(sorted_programs, 2):
        values = [
            prog.university,
            prog.city,
            prog.program,
            prog.language,
            prog.deadline_wise or "",
            prog.deadline_sose or "",
            prog.german_requirement or "",
            prog.english_requirement or "",
            prog.nc_value or "Zulassungsfrei",
            "Evet" if prog.uni_assist_required else "Hayır",
            "Evet" if prog.conditional_admission else "Hayır",
            status_labels.get(prog.eligibility, prog.eligibility),
            prog.eligibility_reason or "",
            prog.url or "",
        ]
        fill = status_fills.get(prog.eligibility)
        for col_idx, val in enumerate(values, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.alignment = Alignment(
                vertical="center",
                wrap_text=(col_idx == 13),   # Değerlendirme sütunu
            )
            if fill:
                cell.fill = fill
        ws.row_dimensions[row_idx].height = 18

    # Özet sayfası
    ws2 = wb.create_sheet("Özet")
    ws2.column_dimensions["A"].width = 22
    ws2.column_dimensions["B"].width = 28

    summary_data = [
        ("Öğrenci",           profile.name),
        ("Hedef Alan",         profile.desired_field),
        ("Derece",             profile.degree_type),
        ("GPA",                profile.gpa_german or profile.gpa_turkish or "—"),
        ("Almanca",            profile.german_level or "Yok"),
        ("İngilizce",          profile.english_level or "Yok"),
        ("",                   ""),
        ("SONUÇLAR",           ""),
        ("Toplam Program",     len(programs)),
        ("✅ Uygun",           sum(1 for p in programs if p.eligibility == "uygun")),
        ("⚠️ Şartlı",         sum(1 for p in programs if p.eligibility == "sartli")),
        ("❌ Uygun Değil",    sum(1 for p in programs if p.eligibility == "uygun_degil")),
        ("",                   ""),
        ("Rapor Tarihi",       datetime.now().strftime("%d.%m.%Y %H:%M")),
        ("Kaynak",             "AES — aes-kompass.com"),
    ]
    for r, (label, value) in enumerate(summary_data, 1):
        a = ws2.cell(row=r, column=1, value=label)
        b = ws2.cell(row=r, column=2, value=value)
        a.font = Font(bold=bool(label and label != ""))
        if label == "SONUÇLAR":
            a.font = Font(bold=True, color="0A1F44", size=11)

    out_path = output_dir / f"universite_listesi_{datetime.now().strftime('%Y%m%d')}.xlsx"
    wb.save(str(out_path))
    print(f"📊 Excel raporu: {out_path}")
    return out_path
